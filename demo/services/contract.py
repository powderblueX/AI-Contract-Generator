import os
import shutil
import docx
from docx import Document
from docx.text.paragraph import Paragraph
import json
from datetime import datetime
from sentence_transformers import SentenceTransformer, util
import torch
import numpy as np
from pathlib import Path
from typing import List, Dict, Tuple, Optional, Union, Any, cast, Sequence
from core.config import *
from dashscope import Generation
from dashscope.api_entities.dashscope_response import GenerationResponse, Message
import re
import gc
import faiss
import time
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PyQt5.QtWidgets import QTextEdit
from io import BytesIO
import uuid
import traceback

class ContractService:
    """
    合同服务核心类
    Contract Service Core Class
    负责合同模板管理、嵌入模型加载、FAISS索引、合同推荐与生成等核心业务逻辑。
    """
    def __init__(self):
        """
        初始化合同服务，加载所有必要资源。
        Initialize contract service, load all required resources.
        """
        self.model: Optional[SentenceTransformer] = None
        self.template_embeddings: Dict[str, np.ndarray] = {}
        self.index: Dict[str, faiss.IndexFlatL2] = {}
        self.filenames: List[str] = []
        self.template_types: Dict[str, List[str]] = {}
        self.desc_edit: Optional[QTextEdit] = None  # 将在UI初始化时设置
        
        # 初始化关键组件
        self.load_model()                  # 加载嵌入模型
        self.load_template_types()         # 加载合同类型映射
        self.load_faiss_index()            # 加载 FAISS 索引
    
    # 加载嵌入模型
    def load_model(self):
        """
        加载SentenceTransformer嵌入模型。
        Load SentenceTransformer embedding model.
        优先本地加载，不存在则自动下载。
        """
        try:
            # 优先从本地加载模型
            model_path = os.path.join(MODEL_DIR, MiniLM_MODEL_NAME)
            # 确保模型目录存在
            os.makedirs(MODEL_DIR, exist_ok=True)
            if os.path.exists(model_path):
                self.model = SentenceTransformer(model_path)
                print(f"从本地加载模型成功: {model_path}")
            else:
                # 本地模型不存在时从HuggingFace下载
                self.model = SentenceTransformer(MiniLM_MODEL_NAME)
                # 保存模型到本地
                self.model.save(model_path)
                print(f"模型下载并保存成功: {model_path}")
        except Exception as e:
            if "MaxRetryError" in str(e) or "Failed to establish a new connection" in str(e):
                print(f"加载模型失败: 无法连接到 Hugging Face。请手动下载模型 '{MiniLM_MODEL_NAME}' 并放置到 {model_path}\n下载地址: https://huggingface.co/sentence-transformers/{MiniLM_MODEL_NAME}")
            else:
                print(f"加载模型失败: {str(e)}")
    
    # 加载 FAISS 索引
    def load_faiss_index(self):
        """
        从磁盘加载 FAISS 索引。
        Load FAISS index from disk.
        """
        try:
            print(RAG_DIR)
            for type_dir in os.listdir(RAG_DIR):
                type_path = os.path.join(RAG_DIR, type_dir)
                if os.path.isdir(type_path):
                    for index in os.listdir(type_path):
                        if index.endswith('.index'):
                            index_path = os.path.join(type_path, index)
                            # 使用 faiss.read_index 加载索引
                            faiss_index = faiss.read_index(index_path)
                            self.index[type_dir] = faiss_index
                            print(f"{type_dir} 的 FAISS 索引已加载，包含 {self.index[type_dir].ntotal} 个向量")
        except Exception as e:
            print(f"从磁盘加载 FAISS 索引失败: {str(e)}")

    # 加载合同类型映射
    def load_template_types(self):
        """加载所有合同类型"""
        try:
            self.template_types = {}
            for type_dir in os.listdir(RAG_DIR):
                type_path = os.path.join(RAG_DIR, type_dir)
                if os.path.isdir(type_path):
                    filenames_path = os.path.join(type_path, "filenames.txt")
                    with open(filenames_path, 'r', encoding='utf-8') as f:
                        templates = f.readlines()
                    templates = [os.path.splitext(line.strip())[0] for line in templates if line.strip()]
                    if templates:
                        self.template_types[type_dir] = templates
        except Exception as e:
            print(f"加载合同类型失败: {str(e)}")
    
    # 获取所有合同类型及其模板
    def get_template_types(self) -> Dict[str, List[str]]:
        """获取所有合同类型及其模板"""
        return self.template_types

    # 获取指定类型的所有模板
    def get_templates_by_type(self, template_type: str) -> List[str]:
        """获取指定类型的所有模板"""
        return self.template_types.get(template_type, [])

    # 读取 txt 文件中的每一行内容
    def read_txt_file(self, file_path):
        with open(file_path, 'r', encoding='utf-8') as file:
            lines = file.readlines()
        # 去除每行末尾的换行符和多余的空白字符
        lines = [os.path.splitext(line.strip())[0] for line in lines if line.strip()]
        return lines

    # 为每一行分配 ID，并生成字典列表
    def create_data_with_ids(self, lines):
        data_list = []
        for idx, line in enumerate(lines, start=1):  # 从 1 开始分配 ID
            data_list.append({
                "id": idx,
                "name": line
            })
        return data_list

    # 加载JSON文件并返回其中的 placeholders 数组
    def load_contract_json(self, file_path):
        """
        从指定路径加载 JSON 文件并返回其内容。
        
        :param file_path: JSON 文件路径
        :return: 加载的 JSON 数据（列表）
        """
        try:
            # 打开 JSON 文件并加载内容
            with open(file_path, "r", encoding="utf-8") as f:
                data = json.load(f) 
            return data["placeholders"] # 一个列表
        except FileNotFoundError:
            print(f"错误：文件 {file_path} 未找到！")
        except json.JSONDecodeError:
            print(f"错误：文件 {file_path} 不是有效的 JSON 格式！")
        except Exception as e:
            print(f"发生未知错误：{e}")
        return []

    # 构造发送给AI模型的提示消息
    def format_message(self, text, ph):
        messages = [
            {'role': 'system', 'content': '你是一个专业的法律文书提取助手，擅长从法律文本中精确提取关键信息。请遵循以下规则：1.严格按照给定关键词提取信息; 2.如找不到信息返回空字符串; 3.只提取实际文本中明确存在的内容，不要推测或创造信息; 4.返回严格的JSON格式，不要添加任何额外说明; 5.确保每个关键词都在结果中，即使值为空字符串。'},
            {'role': 'user', 'content': f'''请从以下文本中提取与关键词相关的具体信息，并以JSON格式返回结果。

    文本内容：
    {text}

    需要提取的关键词：
    {ph}

    提取规则：
    1. 仔细阅读文本，找出与每个关键词精确对应的信息
    2. 如果文本中没有某个关键词的相关信息，请将其值设为空字符串("")
    3. 直接返回JSON格式，不要添加任何解释
    4. 每个关键词必须出现在返回结果中，即使其值为空
    5. 对于明确有否定回答的项目，请提取相应内容（如"否"）而不是留空

    示例1：
    文本：会议时间：2023年10月25日，签订地点：北京市海淀区，合同金额：10万元整，联系人：张三，电话：13800138000
    关键词：会议时间,签订地点,合同金额,甲方代表,乙方代表
    应返回：
    {{{{
    "会议时间": "2023年10月25日",
    "签订地点": "北京市海淀区",
    "合同金额": "10万元整",
    "甲方代表": "",
    "乙方代表": ""
    }}}}

    示例2：
    文本：保管物品：公司重要文件，保管期限：自2023年1月1日至2024年1月1日，保管方式：密封保存，是否需要特殊保管：不需要，保管地点：公司档案室
    关键词：保管物品,保管期限,保管方式,是否需要特殊保管,特殊保管措施内容,保管地点
    应返回：
    {{{{
    "保管物品": "公司重要文件",
    "保管期限": "自2023年1月1日至2024年1月1日",
    "保管方式": "密封保存",
    "是否需要特殊保管": "不需要",
    "特殊保管措施内容": "",
    "保管地点": "公司档案室"
    }}}}

    请确保返回的JSON格式正确，只包含要求的关键词及其对应值，不要添加任何解释或额外信息。'''}]
        return messages

    # 清理AI返回的JSON响应
    def clean_contract_json(self, response_text):
        """
        清理API返回的JSON文本，处理各种可能的格式问题，确保总是返回一个有效的字典
        """
        try:
            # 尝试找出JSON内容
            start = response_text.find('{')
            end = response_text.rfind('}')
            
            if start == -1 or end == -1:
                print(f"警告: 响应中没有找到JSON结构 (start={start}, end={end})")
                # 尝试寻找markdown格式的json
                pattern = r'```(?:json)?\s*\n([\s\S]*?)\n\s*```'
                match = re.search(pattern, response_text)
                if match:
                    potential_json = match.group(1).strip()
                    print(f"从markdown代码块中提取JSON: {potential_json[:100]}...")
                    try:
                        result = json.loads(potential_json)
                        return result
                    except json.JSONDecodeError as e:
                        print(f"从markdown代码块提取的内容不是有效JSON: {str(e)}")
                
                # 如果找不到任何有效JSON，返回空字典
                return {}
            
            # 提取JSON部分
            json_text = response_text[start:end+1].strip()
            
            # 尝试解析JSON
            try:
                cleaned_json = json.loads(json_text)
                return cleaned_json
            except json.JSONDecodeError as e:
                print(f"尝试修复格式错误的JSON: {str(e)}")
                # 尝试修复常见的JSON格式问题
                
                # 1. 处理单引号替换为双引号的问题
                json_text = json_text.replace("'", "\"")
                
                # 2. 处理非法的JSON值（如Python None, True, False）
                json_text = json_text.replace("None", "null")
                json_text = json_text.replace("True", "true")
                json_text = json_text.replace("False", "false")
                
                # 3. 处理末尾可能多出的逗号
                json_text = re.sub(r',\s*}', '}', json_text)
                json_text = re.sub(r',\s*]', ']', json_text)
                
                try:
                    cleaned_json = json.loads(json_text)
                    return cleaned_json
                except json.JSONDecodeError as e:
                    print(f"JSON修复失败: {str(e)}")
                    
                    # 作为最后的尝试，使用更宽松的eval函数
                    try:
                        import ast
                        # 使用ast.literal_eval代替eval，更安全
                        dict_text = json_text.replace("null", "None")
                        dict_text = dict_text.replace("true", "True")
                        dict_text = dict_text.replace("false", "False")
                        cleaned_dict = ast.literal_eval(dict_text)
                        
                        # 确保是字典类型
                        if isinstance(cleaned_dict, dict):
                            print(f"使用ast.literal_eval成功解析")
                            return cleaned_dict
                        else:
                            print(f"ast.literal_eval结果不是字典: {type(cleaned_dict)}")
                            return {}
                    except Exception as e:
                        print(f"所有JSON解析尝试都失败: {str(e)}")
                        return {}
        
        except Exception as e:
            # 捕获所有异常，确保总是返回一个字典
            print(f"clean_contract_json函数出现未处理的异常: {str(e)}")
            return {}

    # 清理合同文档
    def clean_contract(self, docx):
        """
        清理合同中的多余符号，保持格式
        
        Args:
            docx: Document对象
            
        Returns:
            清理后的Document对象
        """
        # 清理段落中的占位符标记
        for paragraph in docx.paragraphs:
            # 使用正则表达式匹配剩余的空占位符标记
            for run in paragraph.runs:
                # 移除空的占位符标记 {xxxx}
                run.text = re.sub(r'\{\s*\}', '', run.text)
                
            # 设置段落对齐方式为向左对齐（如果有明显格式问题）
            # 仅在段落没有明确设置对齐方式时使用
            if paragraph.alignment is None:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        return docx

    # 合同信息提取主流程
    def extract_ph(self, text, ph_json):
        # 从text中提取占位符对应的信息
        # return json格式
        try:
            # 使用类方法加载JSON
            ph = self.load_contract_json(ph_json)
            print(f"从JSON文件加载的占位符: {ph}")
            
            if not ph:
                print(f"警告: 占位符列表为空，返回空字典")
                return {}
                
            ph = ','.join(ph)
            print(f"用户输入文本: {text}")
            print(f"需要提取的关键词: {ph}")

            # 检查输入是否为空
            if not text or not text.strip():
                print("警告: 用户输入为空，返回空字典")
                return {}

            # 记录完整的用户输入到日志，便于调试
            print(f"=== 详细调试信息 - 开始 ===")
            print(f"用户输入原始内容(每行):")
            for i, line in enumerate(text.split('\n')):
                print(f"行{i+1}: {line}")
            print(f"用户输入内容长度: {len(text)}")
            print(f"=== 详细调试信息 - 结束 ===")

            messages = self.format_message(text, ph)

            # 记录传给模型的消息
            print(f"发送给模型的messages:")
            for i, msg in enumerate(messages):
                print(f"消息{i+1} - 角色: {msg['role']}")
                print(f"消息{i+1} - 内容(前100字符): {msg['content'][:100]}...")
            
            # 调用API生成响应
            print(f"开始调用大模型API...")
            try:
                response = Generation.call(
                    api_key=API_KEY,
                    model=LLM_MODEL_NAME,
                    messages=messages,
                    result_format="message"
                )
                print(f"大模型API调用完成")
            except Exception as e:
                print(f"大模型API调用失败: {str(e)}")
                return {}  # API调用失败时返回空字典
            
            response_text = ""
            if hasattr(response, 'output') and hasattr(response.output, 'choices'):
                # 获取处理后的结果
                try:
                    response_text = response['output']['choices'][0]['message']['content']
                    print(f"模型返回的原始响应: {response_text}")
                except (KeyError, IndexError, TypeError) as e:
                    print(f"解析模型响应失败: {str(e)}")
                    return {}  # 解析响应失败时返回空字典
            else:
                print("API响应格式不正确")
                print(f"响应对象内容: {str(response)}")
                return {}  # 响应格式不正确时返回空字典
            
            print(f"开始解析JSON响应...")
            response_json = self.clean_contract_json(response_text)
            print(f"解析后的JSON数据: {response_json}")
            
            # 检查解析结果是否为空
            if not response_json:
                print(f"警告: 解析后的JSON为空! 这可能导致合同填充失败")
                return {}  # 解析结果为空时返回空字典
            return response_json
        except Exception as e:
            print(f"提取占位符过程中发生异常: {str(e)}")
            # 捕获所有异常，确保函数总是返回字典而不是抛出异常
            return {}

    # 从用户输入提取合同关键词
    def extract_keywords(self, text):
        """
        从用户输入中提取合同相关的关键词
        增强版：提取更多法律专业术语，增加合同类型识别能力
        """
        messages = [
                    { 'role': 'system', 'content': '''
                        # 角色定位
                        您是中国合同法律框架分析专家，专精合同核心要素提取。请严格遵循：
                        1. 聚焦合同法律关系和权利义务框架
                        2. 排除具体数值（金额/日期等）和个性化条款
                        3. 采用《民法典》合同编术语体系''' },
                    { 'role': 'user', 'content': f'''
                        ## 合同要素提取指令
                        请从文本中提取合同相关的专业关键词，并按以下规则处理：
                     
                        ### 核心提取维度
                        1. **合同性质要素**
                        - 合同类型（如"保管合同"、"租赁合同"）
                        - 法律特征（如"有偿/无偿"、"要式/不要式"）

                        2. **主体要素**
                        - 当事人法律身份（如"保管人"、"寄存人"）
                        - 主体资质要求（如"特殊经营资质"）

                        3. **客体要素**
                        - 标的物法律属性（如"动产/不动产"、"种类物/特定物"）
                        - 标的特殊性质（如"危险品"、"贵重物"）

                        4. **权利义务要素**
                        - 核心义务（如"验收义务"、"保管责任"）
                        - 特殊权利（如"留置权"、"转授权"）
                        - 责任机制（如"违约责任"）

                        5. **程序要素**
                        - 成立要件（如"书面形式"）
                        - 履行程序（如"交付验收"）
                        - 争议解决机制

                        ### 处理规范
                        - **重点提取**：反映合同法律本质的抽象概念
                        - **术语要求**：
                            * 使用《民法典》标准术语（如"瑕疵担保"而非"质量问题"）
                            * 保持概念完整性（如"保管物转交"而非简单"转交"）
                            * 地理位置包含行政区划，且必须包含对应的省级行政区名（如"福建省"、"上海市"）
                        - **必须排除**：具体金额、日期、模糊表述（如"八十万人民币"、"2022年6月7日"、"大概"、"可能"等）
                        - **排除范围**：非法律关系的程序性描述（如"签字盖章"）
                     
                        ## 输出要求
                        1. 严格按重要性降序排列（法律效力最高者在前）
                        2. 关键词间用逗号分隔，不加序号
                        3. 不包含任何解释性文字
                        4. 使用简体中文专业术语
                        5. 确保识别出尽可能多的专业词汇

                        ## 示例说明
                        输入：我需要与某公司就一处位于北京海淀区的商业房产签订为期5年的租赁合同，月租金约为2万元。
                        输出：商业房产租赁合同, 商业房产, 房产, 北京, 租期, 月租金, 商业地产, 租赁权, 房屋租赁, 租赁

                        ## 待分析文本
                        {text}
                        '''}
                ]
        
        # 调用大模型API生成响应
        response = Generation.call(
            api_key=API_KEY,
            model=LLM_MODEL_NAME,
            messages=messages,
            result_format="message"
        )
        
        response_text = ""
        if hasattr(response, 'output') and hasattr(response.output, 'choices'):
            response_text = response['output']['choices'][0]['message']['content']
            print(f"提取的合同关键词: {response_text}")
        else:
            print("API响应格式不正确")

        # 优化文本清理：保留法律术语中可能包含的特殊字符
        return self.clean_text_for_legal(response_text)

    # 法律文本专用清理
    def clean_text_for_legal(self, text):
        """
        专为法律文本优化的清理函数，保留更多法律术语的特征
        """
        # 去除多余空白字符
        text = re.sub(r"\s+", " ", text).strip()
        
        # 保留一些法律文本中重要的标点符号（如括号、引号等）
        # 仅删除顿号、逗号等作为分隔符的符号
        text = re.sub(r"[、，,]", " ", text)
        
        # 其他特殊字符保留，不做过度清理
        
        return text.strip()

    # 模板填充核心功能
    def fill_template(self, template_docx, ph_json):
        """
        把占位符对应的信息填入模版的占位符，并保持原始格式
        
        Args:
            template_docx: Document对象，加载的docx模板
            ph_json: 字典，包含占位符名称和对应的值
            
        Returns:
            填充后的Document对象
        """
        
        # 检查输入参数的有效性
        if not ph_json or not isinstance(ph_json, dict):
            print("警告: 占位符JSON为空或不是字典格式，将使用空字典")
            ph_json = {}
        
        # 首先收集所有需要替换的占位符
        all_placeholders = {f"{{{key}}}": value for key, value in ph_json.items()}
        
        try:
            # 遍历文档中的所有段落
            for paragraph in template_docx.paragraphs:
                # 存储所有runs的文本内容，以便后续处理
                runs_text = []
                for run in paragraph.runs:
                    runs_text.append(run.text)
                
                # 检查段落中是否包含占位符
                paragraph_has_placeholder = False
                for ph in all_placeholders.keys():
                    if any(ph in run_text for run_text in runs_text):
                        paragraph_has_placeholder = True
                        break
                        
                if paragraph_has_placeholder:            
                    # 使用run级别替换，保留格式
                    for i, run in enumerate(paragraph.runs):
                        # 检查当前run是否包含占位符
                        original_text = run.text
                        modified_text = original_text
                        
                        # 替换所有匹配的占位符
                        for ph, value in all_placeholders.items():
                            if ph in modified_text:
                                # 确保value是字符串类型
                                if value is None:
                                    value = ""
                                if not isinstance(value, str):
                                    value = str(value)
                                    
                                # 使用空格替代"没有内容"
                                replacement = "    " if value == "没有内容" else value
                                modified_text = modified_text.replace(ph, replacement)
                        
                        # 只有当文本发生变化时才更新
                        if modified_text != original_text:
                            run.text = modified_text
            
            # 再次检查是否还有未替换的占位符
            for paragraph in template_docx.paragraphs:
                for run in paragraph.runs:
                    # 安全检查，确保run.text是字符串
                    if not hasattr(run, 'text') or not isinstance(run.text, str):
                        continue
                        
                    if '{' in run.text and '}' in run.text:
                        pattern = r'\{[^{}]+\}'  # 匹配{xxx}格式
                        matches = re.findall(pattern, run.text)
                        if matches:
                            print(f"警告：发现未替换的占位符: {matches}")
                            # 将未替换的占位符替换为空格，而不是保留
                            for match in matches:
                                run.text = run.text.replace(match, "    ")
            
            # 检查表格中的占位符
            for table in template_docx.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            # 存储所有runs的文本内容
                            runs_text = []
                            for run in paragraph.runs:
                                # 安全检查
                                if hasattr(run, 'text') and isinstance(run.text, str):
                                    runs_text.append(run.text)
                            
                            # 检查段落中是否包含占位符
                            paragraph_has_placeholder = False
                            for ph in all_placeholders.keys():
                                if any(ph in run_text for run_text in runs_text):
                                    paragraph_has_placeholder = True
                                    break
                                    
                            if paragraph_has_placeholder:
                                print(f"在表格单元格段落中发现占位符: {paragraph.text}")
                                
                                # 使用run级别替换，保留格式
                                for i, run in enumerate(paragraph.runs):
                                    # 检查当前run是否包含占位符
                                    # 安全检查
                                    if not hasattr(run, 'text') or not isinstance(run.text, str):
                                        continue
                                        
                                    original_text = run.text
                                    modified_text = original_text
                                    
                                    # 替换所有匹配的占位符
                                    for ph, value in all_placeholders.items():
                                        if ph in modified_text:
                                            # 确保value是字符串类型
                                            if value is None:
                                                value = ""
                                            if not isinstance(value, str):
                                                value = str(value)
                                                
                                            # 使用空格替代"没有内容"
                                            replacement = "    " if value == "没有内容" else value
                                            modified_text = modified_text.replace(ph, replacement)
                                    
                                    # 只有当文本发生变化时才更新
                                    if modified_text != original_text:
                                        run.text = modified_text
            
            # 处理表格中可能剩余的未替换占位符
            for table in template_docx.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                # 安全检查
                                if not hasattr(run, 'text') or not isinstance(run.text, str):
                                    continue
                                    
                                if '{' in run.text and '}' in run.text:
                                    pattern = r'\{[^{}]+\}'  # 匹配{xxx}格式
                                    matches = re.findall(pattern, run.text)
                                    if matches:
                                        print(f"警告：表格中发现未替换的占位符: {matches}")
                                        # 将未替换的占位符替换为空格，而不是保留
                                        for match in matches:
                                            run.text = run.text.replace(match, "    ")
            
            # 清理空白行和多余符号
            template_docx = self.clean_contract(template_docx)
            
            return template_docx
        except Exception as e:
            print(f"填充模板时发生异常: {str(e)}")
            # 返回原始模板，而不是抛出异常，确保流程能继续
            return template_docx

    # 将填充后的文档转为内存字节流
    def template_to_bytes(self, template_docx):
        # 将文档保存到内存中的字节流
        byte_stream = BytesIO()
        template_docx.save(byte_stream)  # 将文档保存到字节流
        byte_stream.seek(0)  # 移动指针到字节流的开头

        return byte_stream

    # 加载所有 .txt 文件的内容
    def load_txt_files(self, folder_path):
        """
        从指定文件夹中加载所有 .txt 文件的内容。
        返回一个包含 (文件名, 文件内容) 的列表。
        """
        documents = []
        for filename in os.listdir(folder_path):
            if filename.endswith(".txt"):
                file_path = os.path.join(folder_path, filename)
                with open(file_path, "r", encoding="utf-8") as f:
                    content = f.read().strip()
                    if content:  # 忽略空文件
                        documents.append((filename, content))
        return documents

    # 通用文本清理
    def clean_text(self, text):
        """
        清理文本内容：
        - 去除多余空白字符（包括换行符、制表符等）。
        - 去除顿号、逗号等分隔符。
        - 去除特殊字符（如标点符号）。
        :param text: 原始文本
        :return: 清理后的文本
        """
        # 去除多余空白字符
        text = re.sub(r"\s+", " ", text).strip()
        
        # 去除顿号、逗号等分隔符
        text = re.sub(r"[、，,]", " ", text)
        
        # 去除其他特殊字符（可选）
        text = re.sub(r"[^\w\s]", "", text)
        
        return text.strip()

    # 将文本内容向量化
    def generate_embeddings(self, documents, model):
        """
        使用 SentenceTransformer 模型将文档内容转换为向量。
        返回文件名列表和对应的向量矩阵。
        """
        filenames, contents = zip(*documents)  # 分离文件名和内容
        # 清理内容
        contents = [self.clean_text(content) for content in contents]

        embeddings = model.encode(contents, convert_to_tensor=False)  # 生成向量
        return list(filenames), np.array(embeddings).astype('float32')

    # 构建 FAISS 索引并保存
    def build_faiss_index(self, embeddings, dimension):
        """
        使用 FAISS 构建向量索引。
        """
        index = faiss.IndexFlatL2(dimension)  # 使用 L2 距离
        index.add(embeddings)  # 添加向量到索引
        return index

    # 查询知识库
    def search_in_knowledge_base(self, query, top_k=5):
        """
        根据查询字符串在知识库中搜索最相似的文档。
        """

        """使用预加载资源进行知识库搜索"""
        if not self.model or not self.index:
            print("模型或索引未初始化")
            return []
    
        try:
            # 使用预加载的模型生成查询向量
            query_vector = self.model.encode([query], convert_to_tensor=False).astype('float32')
            
            # 使用预加载的FAISS索引进行搜索
            distances, indices = self.index.search(query_vector, top_k)
            
            # 使用预加载的文件名列表获取结果
            return [(self.filenames[i], distances[0][j]) for j, i in enumerate(indices[0])]
        except Exception as e:
            print(f"搜索失败: {str(e)}")
            return []

    # 构造模板关键词提取提示
    def format_template_keywords(self, docx_content):
        messages = [
            {'role': 'system', 'content': 'You are a legal expert specialized in contract law.'},
            {'role': 'user', 'content': f'''请分析以下合同模板内容，提取能全面表示该合同特征的关键词。作为法律专家，需要提取以下五类关键词：

    1. 合同类型分类：识别该合同属于哪类合同（如买卖类、租赁类、服务类等）
    2. 适用场景：该合同特别适用的具体场景（如商品房买卖、二手车交易等）
    3. 关键条款：该合同中特有的核心条款（如质量验收标准、交付方式、价款计算等）
    4. 法律术语：合同中出现的专业法律术语
    5. 行业术语：特定行业相关术语

    请确保每一类都尽量提取，用逗号分隔词语，各类之间用分号(;)分隔。不要提取通用合同中都有的内容（如违约责任、争议解决等）。

    合同模板内容：
    {docx_content}

    输出格式示例：
    买卖合同,不动产买卖;商品房,期房,现房,商业地产;首付款,银行按揭,产权交割,面积差异处理;买卖标的,交付使用,权属转移;房地产开发,预售许可证
    '''}
        ]
        return messages

    # 从合同模板提取结构化关键词（分5类）
    def extract_template_keywords(self, docx_content):
        """
        从合同模板中提取关键词，增强版采用分类提取方式
        """

        """使用预定义的消息格式提取模板关键词"""
        try:
            # 构造消息（可复用预定义的格式）
            messages = self.format_template_keywords(docx_content)
            
            # 调用API（API密钥可考虑在初始化时加载）
            response = Generation.call(
                api_key=API_KEY,
                model=LLM_MODEL_NAME,
                messages=messages,
                result_format="message"
            )

            response_text = ""
            if hasattr(response, 'output') and hasattr(response.output, 'choices'):
                # 获取处理后的结果
                response_text = response['output']['choices'][0]['message']['content']
                print(f"提取的合同模板结构化关键词: {response_text}")
            else:
                print("API响应格式不正确")
        
            return response_text
        except Exception as e:
            print(f"关键词提取失败: {str(e)}")
            return ""

    # 批量处理合同模板
    def process_contract_templates(self, input_folder, output_folder):
        """
        处理合同模板文件夹中的所有docx文件，提取关键词并保存
        同时生成一个合同类型分类映射文件，用于提升检索准确性
        """

        """使用预加载模型处理合同模板"""
        if not self.model:
            print("模型未加载，无法处理模板")
            return {}

        # 确保输出文件夹存在
        os.makedirs(output_folder, exist_ok=True)
        
        # 获取所有docx文件
        all_files = [f for f in os.listdir(input_folder) if f.endswith('.docx')]
        total_files = len(all_files)
        
        # 创建合同类型映射
        contract_categories = {}
        
        for index, file_name in enumerate(all_files, 1):
            try:
                print(f"正在处理文件: {file_name} 进度:{index}/{total_files}")
                
                # 读取docx文件
                file_path = os.path.join(input_folder, file_name)
                docx_content = self.get_docx(file_path)
                
                # 提取关键词
                keywords = self.extract_template_keywords(docx_content)
                
                # 保存到txt文件
                txt_path = os.path.join(output_folder, file_name.replace('.docx', '.txt'))
                self.save_string_to_txt(keywords, txt_path)
                
                # 提取并记录合同类型分类
                try:
                    categories = keywords.split(';')[0].strip()
                    contract_categories[file_name.replace('.docx', '')] = categories
                except:
                    print(f"无法从 {file_name} 提取合同类型")
                
            except Exception as e:
                print(f"处理文件 {file_name} 时出错: {e}")
                continue
        
        # 保存合同类型映射
        category_path = os.path.join(output_folder, "_contract_categories.json")
        try:
            with open(category_path, 'w', encoding='utf-8') as f:
                json.dump(contract_categories, f, ensure_ascii=False, indent=2)
            print(f"合同类型映射已保存至 {category_path}")
        except Exception as e:
            print(f"保存合同类型映射失败: {e}")
            
        return contract_categories

    # 深度分析用户需求
    def analyze_user_needs(self, text):
        """
        分析用户需求，提取更深层次的意图和偏好，
        加强对无关内容的识别
        """
        # 检查输入文本长度，如果过短则直接返回无相关合同
        if not text or len(text.strip()) < 8:
            return {
                "contract_category": "无相关合同",
                "specific_type": "N/A",
                "special_concerns": []
            }
        
        # 快速检查是否包含常见问候语或无关内容
        greeting_patterns = [
            r'^你好[！!.,，。\s]*$', 
            r'^您好[！!.,，。\s]*$',
            r'^早上好[！!.,，。\s]*$',
            r'^下午好[！!.,，。\s]*$',
            r'^晚上好[！!.,，。\s]*$',
            r'^嗨[！!.,，。\s]*$',
            r'^hi[！!.,，。\s]*$',
            r'^hello[！!.,，。\s]*$'
        ]
        
        for pattern in greeting_patterns:
            if re.match(pattern, text.strip(), re.IGNORECASE):
                return {
                    "contract_category": "无相关合同",
                    "specific_type": "N/A",
                    "special_concerns": []
                }
        
        messages = [{ 'role': 'system', 'content': '''
                        # 角色指令
                        您作为中国律所资深合同分析师，需从对话文本中精准提取合同需求。请严格遵循以下规则：
                        1. 仅当文本明确涉及合同条款、权利义务或商业安排时，才识别合同类型
                        2. 无关内容（问候/闲聊/非合同事务）一律返回"无相关合同"
                        3. 特殊关注点需直接引用文本关键词，不做主观扩展''' }, 
                     { 'role': 'user', 'content': f'''
                        ## 合同需求分析任务
                        请从以下维度解析用户输入：
                        ### 分析维度
                        1. **合同大类**（如：服务类、买卖类、租赁类、合伙类、雇佣类等）
                        2. **具体类型**（如：软件开发合同、设备租赁合同、股权投资协议等）
                        3. **特殊关注点**（仅提取文本明确提及的敏感条款，如："付款期限要明确"→"付款期限"）

                        ## 处理原则
                        - 模糊表述（如"我们需要合作"）视为无相关合同
                        - 特殊关注点必须来源于文本直接表述
                        - 排除非约束性表述（如"希望长期合作"）

                        ## 输出要求
                        请严格生成JSON对象，包含且仅包含以下字段：
                        {{
                            "contract_category": "大类名称 或 '无相关合同'",
                            "specific_type": "具体类型名称 或 'N/A'",
                            "special_concerns": ["关键词1", "关键词2"]  // 无则留空数组
                        }}

                        ## 用户输入文本
                        {text}'''}
                    ]
        
        # 调用API生成响应
        response = Generation.call(
            api_key=API_KEY,
            model=LLM_MODEL_NAME,
            messages=messages,
            result_format="message"
        )
        
        response_text = ""
        if hasattr(response, 'output') and hasattr(response.output, 'choices'):
            # 获取处理后的结果
            response_text = response['output']['choices'][0]['message']['content']
            print(f"用户需求分析结果: {response_text}")
            
            # 尝试解析JSON
            try:
                # 清理JSON文本
                if response_text.startswith("```json"):
                    response_text = response_text[7:]
                if response_text.endswith("```"):
                    response_text = response_text[:-3]
                    
                analysis = json.loads(response_text.strip())
                
                # 确保category和type都是标准格式
                if not analysis["contract_category"] or analysis["contract_category"].lower() in ["无", "无关", "无关合同", "不相关", "不适用"]:
                    analysis["contract_category"] = "无相关合同"
                    
                if not analysis["specific_type"] or analysis["specific_type"].lower() in ["无", "无关", "不相关", "不适用", "n/a", "na"]:
                    analysis["specific_type"] = "N/A"
                    
                return analysis
            except:
                print("JSON解析失败，返回原始文本")
                return {
                    "contract_category": "无相关合同",
                    "specific_type": "N/A",
                    "special_concerns": []
                }
        else:
            print("API响应格式不正确")
            return {
                "contract_category": "无相关合同",
                "specific_type": "N/A",
                "special_concerns": []
            }

    # 加载合同分类映射
    def load_contract_categories(self, template_dir, contract_type):
        """
        加载合同类型分类映射，如果不存在_contract_categories.json，则直接读取txt关键词文件
        """
        try:
            # 首先尝试加载JSON文件（如果存在）
            category_path = os.path.join(template_dir, contract_type, "关键词", "_contract_categories.json")
            if os.path.exists(category_path):
                try:
                    with open(category_path, 'r', encoding='utf-8') as f:
                        return json.load(f)
                except Exception as e:
                    print(f"加载合同类型分类映射失败: {e}")
            
            # 如果JSON文件不存在，则遍历关键词文件夹中的所有txt文件
            print("_contract_categories.json不存在，尝试从关键词txt文件构建映射")
            categories_map = {}
            
            # 确保路径存在
            keywords_folder = os.path.join(template_dir, contract_type, "关键词")
            if not os.path.exists(keywords_folder):
                print(f"关键词文件夹不存在: {keywords_folder}")
                return {}
            
            # 遍历文件夹中的所有txt文件
            try:
                for filename in os.listdir(keywords_folder):
                    if filename.endswith(".txt") and not filename.startswith("_"):
                        try:
                            file_path = os.path.join(keywords_folder, filename)
                            with open(file_path, 'r', encoding='utf-8') as f:
                                content = f.read().strip()
                            
                            # 使用文件名（不包括扩展名）作为key，文件内容作为value
                            file_key = os.path.splitext(filename)[0]
                            
                            # 如果内容包含分号(;)，获取第一部分作为合同分类
                            if ";" in content:
                                categories_map[file_key] = content.split(";")[0].strip()
                            else:
                                # 否则使用整个内容
                                categories_map[file_key] = content
                                
                        except Exception as e:
                            print(f"读取关键词文件 {filename} 出错: {str(e)}")
                            continue
                
                print(f"成功从关键词文件构建映射，共 {len(categories_map)} 个合同类型")
                return categories_map
            except Exception as e:
                print(f"遍历关键词文件夹时出错: {str(e)}")
                return {}
        except Exception as e:
            print(f"load_contract_categories函数出现未处理的异常: {str(e)}")
            # 返回空字典而不是抛出异常
            return {}

    # 通过预定义关键词映射识别合同类型
    def extract_contract_type_from_keywords(self, keywords_text):
        """
        从用户输入的关键词中识别可能的合同类型
        """
        # 常见合同类型关键词映射
        contract_type_keywords = {
            "买卖": ["买卖", "购买", "出售", "销售", "采购", "交易"],
            "租赁": ["租赁", "出租", "承租", "房屋租赁", "租用", "租借"],
            "服务": ["服务", "咨询", "顾问", "中介", "代理", "委托"],
            "工程": ["工程", "建设", "施工", "装修", "建筑", "设计", "监理"],
            "劳务": ["劳务", "用工", "雇佣", "聘用", "招聘", "人力资源"],
            "保管": ["保管", "仓储", "存储", "寄存", "托管"],
            "运输": ["运输", "物流", "快递", "货运", "配送", "运送"],
            "许可": ["特许", "许可", "授权", "专利", "商标", "著作权", "知识产权"],
            "融资": ["融资", "借款", "贷款", "抵押", "担保", "质押", "保证"],
            "赠与": ["赠与", "捐赠", "馈赠", "无偿"],
            "修缮": ["修缮", "修理", "维修", "改造", "翻新", "维护", "保养"],
            "承揽": ["承揽", "承接", "接受委托", "完成工作"],
            "加工": ["加工", "定制", "生产制造", "原料加工"],
            "定作": ["定作", "定制", "定做", "定制品", "特制"],
            "房产": ["房产", "房屋", "住宅", "商品房", "不动产"],
            "汽车": ["汽车", "车辆", "二手车", "机动车"],
            "保险": ["保险", "投保", "理赔", "保额"]
        }
        
        identified_types = []
        
        # 查找关键词中包含的合同类型
        for type_name, keywords in contract_type_keywords.items():
            for keyword in keywords:
                if keyword in keywords_text:
                    identified_types.append(type_name)
                    break
        
        return list(set(identified_types))  # 去重

    # 高级检索流程
    def advanced_search_in_knowledge_base(self, query, model, index, filenames, categories_map, user_text, top_k=10):
        """
        高级知识库搜索，结合向量检索和合同类型分类过滤
        
        参数:
        - query: 查询文本（已处理的关键词）
        - model: SentenceTransformer模型
        - index: FAISS索引
        - filenames: 文件名列表
        - categories_map: 合同类型映射
        - user_text: 用户原始输入
        - top_k: 返回结果数量
        
        返回:
        - 排序后的结果列表：[(filename, score), ...]
        """
        try:
            print(f"高级搜索开始，查询文本: {query}")
            print(f"索引中的文件数量: {len(filenames)}")
            print(f"Faiss索引大小: {index.ntotal}")
            
            # 1. 基础向量搜索 (检索更多结果用于后处理)
            query_vector = model.encode([query], convert_to_tensor=False).astype('float32')
            print(f"查询向量生成完成，维度: {query_vector.shape}")
            
            # 检查索引和文件名长度匹配
            if index.ntotal != len(filenames):
                print(f"警告: 索引大小({index.ntotal})与文件数量({len(filenames)})不匹配")
            
            # 计算实际检索数量
            actual_top_k = min(top_k * 2, len(filenames))
            print(f"将检索 {actual_top_k} 个结果")
            
            distances, indices = index.search(query_vector, actual_top_k)
            print(f"检索完成，获得 {len(indices[0])} 个结果")
            
            # 安全检查索引值有效性
            valid_indices = [i for i in indices[0] if 0 <= i < len(filenames)]
            if len(valid_indices) < len(indices[0]):
                print(f"警告: 检索到 {len(indices[0]) - len(valid_indices)} 个无效索引，已过滤")
            
            basic_results = []
            for j, i in enumerate(indices[0]):
                if 0 <= i < len(filenames):
                    basic_results.append((filenames[i], float(distances[0][j])))
                else:
                    print(f"警告: 跳过无效索引 {i}，超出范围 [0, {len(filenames)-1}]")
            
            print(f"基础检索结果数量: {len(basic_results)}")
            
            # 2. 识别用户输入中的合同类型
            contract_types = self.extract_contract_type_from_keywords(user_text + " " + query)
            print(f"从用户输入中识别的合同类型: {contract_types}")
            
            # 3. 根据合同类型进行结果重排序
            final_results = []
            category_matched = []
            category_unmatched = []
            
            # 如果没有识别出合同类型，直接返回基础检索结果
            if not contract_types:
                print("未识别出合同类型，直接返回基础检索结果")
                return basic_results[:top_k]
            
            # 根据合同类型对结果进行分类
            for filename, score in basic_results:
                try:
                    file_key = os.path.splitext(filename.strip())[0]
                    
                    # 查找该文件的合同类型
                    file_categories = categories_map.get(file_key, "")
                    
                    # 判断是否匹配用户需求的合同类型
                    matched = False
                    for ct in contract_types:
                        if ct in file_categories:
                            matched = True
                            # 匹配合同类型的文件得分提升
                            category_matched.append((filename, score * 0.7))  # 降低距离分数，提高排名
                            break
                    
                    if not matched:
                        category_unmatched.append((filename, score))
                except Exception as e:
                    print(f"处理文件 {filename} 时出错: {str(e)}")
                    # 发生错误但不中断，继续处理其他文件
                    continue
            
            print(f"类型匹配结果数量: {len(category_matched)}")
            print(f"类型不匹配结果数量: {len(category_unmatched)}")
            
            # 4. 分别对匹配和未匹配的结果进行排序
            category_matched.sort(key=lambda x: x[1])
            category_unmatched.sort(key=lambda x: x[1])
            
            # 5. 组合结果：优先返回类型匹配的结果，然后是其他结果
            final_results = category_matched + category_unmatched
            
            print(f"最终返回结果数量: {min(top_k, len(final_results))}")
            return final_results[:top_k]
        
        except Exception as e:
            import traceback
            error_details = traceback.format_exc()
            print(f"高级搜索发生异常: {str(e)}")
            print(f"详细错误信息: {error_details}")
            
            # 返回空结果而不是抛出异常
            return []